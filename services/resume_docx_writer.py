"""
Renders TailoredResumeContent into a clean, ATS-friendly .docx using
python-docx (a runtime dependency of this backend already, unlike the
Node docx-js toolchain used for one-off document authoring elsewhere --
this needs to run unattended inside the API process).

Formatting choices are deliberately plain: single column, standard
heading styles, no text boxes/tables-as-layout, no headers/footers with
critical content -- all things that break ATS parsers. "Clean and
recruiter-friendly" per the agreed spec means boring on purpose here.
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from services.resume_tailoring import TailoredResumeContent


def _slugify(text: str) -> str:
    lowered = re.sub(r"[^a-z0-9\s_-]", "", text.strip().lower())
    return re.sub(r"[\s-]+", "_", lowered).strip("_")


def build_resume_filename(candidate_full_name: str, target_title: str) -> str:
    """'Raheel Ahmed Khan' + 'Lead Full Stack Developer' ->
    'raheel_ahmed_khan_lead_full_stack_developer.docx' -- name and
    target role both visible in the filename, per the agreed spec."""
    return f"{_slugify(candidate_full_name)}_{_slugify(target_title)}.docx"


def render_resume_docx(
    content: TailoredResumeContent,
    candidate_full_name: str,
    contact_line: str,
    location: str,
) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Header block
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(candidate_full_name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(content.target_title)
    title_run.font.size = Pt(12)
    title_run.italic = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_para = doc.add_paragraph()
    contact_run = contact_para.add_run(f"{contact_line} | {location}")
    contact_run.font.size = Pt(9.5)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    def add_section_heading(text: str) -> None:
        heading = doc.add_paragraph()
        run = heading.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11.5)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        # simple bottom border for a section rule, without using a table
        pPr = heading._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Professional summary
    if content.professional_summary:
        add_section_heading("Professional Summary")
        doc.add_paragraph(content.professional_summary)

    # Core skills
    if content.core_skills_section:
        add_section_heading("Technical Skills")
        skills_para = doc.add_paragraph(" | ".join(content.core_skills_section))
        skills_para.paragraph_format.space_after = Pt(4)

    # Experience
    if content.experience:
        add_section_heading("Professional Experience")
        for entry in content.experience:
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(f"{entry.get('title', '')} — {entry.get('company', '')}")
            role_run.bold = True
            role_run.font.size = Pt(10.5)

            dates_para = doc.add_paragraph()
            dates_run = dates_para.add_run(entry.get("dates", ""))
            dates_run.italic = True
            dates_run.font.size = Pt(9.5)
            dates_para.paragraph_format.space_after = Pt(2)

            for bullet in entry.get("bullets", []):
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(2)

    # Additional technical exposure -- kept visibly separate from core
    # skills/experience, per the agreed provenance split.
    if content.additional_technical_exposure:
        add_section_heading("Additional Technical Exposure")
        exposure_para = doc.add_paragraph(
            "Working knowledge / self-study, not production experience: "
            + ", ".join(content.additional_technical_exposure)
        )
        exposure_para.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
