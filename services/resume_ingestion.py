"""
Resume ingestion pipeline: base resume file -> extracted text -> Claude
API call that suggests skills + provenance tiers -> SkillInventoryItem
rows with status='pending', queued for admin approval.

This module NEVER writes status='approved' -- that's the one invariant
the whole matching pipeline depends on. A resume update (e.g. a
candidate quietly adding "C++" to an old project because C++ jobs are
trending) can only ever land here as a pending suggestion, never as a
usable skill, until a human reviews it via the approval-queue API.

No network call happens at import time or module load -- the Anthropic
client is constructed lazily inside extract_skills_from_resume_text so
importing this module (e.g. for tests) never requires ANTHROPIC_API_KEY
to be set.
"""
from __future__ import annotations

import io
import json
import os
from dataclasses import dataclass

from sqlalchemy.orm import Session

from db.models import Candidate, ResumeIngestionRun, SkillInventoryItem


@dataclass
class SuggestedSkill:
    skill_name: str
    tier: str  # core | component | secondary | exposure
    source_bullet: str | None
    source_project_or_role: str | None
    confidence: float


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Pulls plain text out of a .docx: paragraphs and table cells, in
    document order. Deliberately simple -- this feeds an LLM extraction
    step next, which tolerates messy/unstructured text far better than a
    strict parser would need to."""
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


_EXTRACTION_SYSTEM_PROMPT = """\
You extract skill claims from a resume's raw text for a recruiting tool. \
For every distinct technology/skill/tool mentioned, output one entry with:
- skill_name: the skill as it should be searched/matched (e.g. "Kafka", "TypeScript", "Terraform")
- tier: one of "core", "component", "secondary", "exposure"
    - core: primary, years of hands-on work, central to the person's role
    - component: clearly an inherent part of something else they have deep experience in
      (e.g. TypeScript mentioned only via Angular work, S3 mentioned only via general AWS work)
    - secondary: real hands-on work described, but clearly supporting/adjacent to their
      primary role (e.g. "supported DevOps team with Terraform configs")
    - exposure: only a course, certification, brief mention, or single small project --
      not demonstrated production work
- source_bullet: the exact resume line/snippet this was inferred from (verbatim, do not paraphrase)
- source_project_or_role: the job/project heading this bullet appeared under, if identifiable, else null
- confidence: your confidence in the tier assignment, 0.0-1.0

Rules:
- Do NOT invent skills that are not mentioned anywhere in the text.
- Do NOT upgrade a tier beyond what the text actually supports -- when in doubt, tier down
  (prefer "exposure" over "secondary", "secondary" over "component"/"core").
- Every entry must be traceable to a specific source_bullet actually present in the input text.
- Output ONLY a JSON array of objects with the fields above. No prose, no markdown fences.
"""


def extract_skills_from_resume_text(resume_text: str) -> list[SuggestedSkill]:
    """Calls the Claude API to produce skill/tier suggestions from raw
    resume text. Suggestions only -- callers must not treat this output
    as approved; it's meant to populate SkillInventoryItem rows with
    status='pending' for human review.

    Raises RuntimeError if ANTHROPIC_API_KEY isn't set, rather than
    silently skipping extraction (a silent skip here would mean a
    resume change never reaches the approval queue at all)."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError(
            "ANTHROPIC_API_KEY is not set -- required for resume skill "
            "extraction. Set it in your .env before running ingestion."
        )

    import anthropic

    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        system=_EXTRACTION_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": resume_text}],
    )

    raw_text = "".join(block.text for block in response.content if block.type == "text")
    raw_text = raw_text.strip()
    if raw_text.startswith("```"):
        raw_text = raw_text.strip("`")
        if raw_text.startswith("json"):
            raw_text = raw_text[len("json"):]
        raw_text = raw_text.strip()

    try:
        parsed = json.loads(raw_text)
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Skill extraction returned non-JSON output: {e}\nRaw: {raw_text[:500]}") from e

    suggestions: list[SuggestedSkill] = []
    for item in parsed:
        try:
            suggestions.append(
                SuggestedSkill(
                    skill_name=str(item["skill_name"]).strip(),
                    tier=str(item["tier"]).strip().lower(),
                    source_bullet=item.get("source_bullet"),
                    source_project_or_role=item.get("source_project_or_role"),
                    confidence=float(item.get("confidence", 0.5)),
                )
            )
        except (KeyError, TypeError, ValueError):
            continue  # skip malformed entries rather than fail the whole batch

    valid_tiers = {"core", "component", "secondary", "exposure"}
    return [s for s in suggestions if s.tier in valid_tiers and s.skill_name]


def ingest_resume(
    session: Session,
    candidate: Candidate,
    resume_file_path: str,
    file_bytes: bytes,
    file_hash: str,
    triggered_by: str = "file_watcher",
    pending_storage_key: str | None = None,
) -> ResumeIngestionRun:
    """Full pipeline for one resume file: extract text -> call Claude ->
    write pending SkillInventoryItem rows -> mark the run completed (or
    failed). Never touches SkillInventoryItem.status='approved'.

    The resume FILE itself has its own, separate approval gate:
    - triggered_by='candidate_upload': the file was saved to
      pending_storage_key, NOT the live resume path -- it only becomes
      the active/tailoring-eligible resume once an admin explicitly
      approves it (see api/routers/candidate_documents.py's resume
      approval endpoints). resume_approval_status stays 'pending'.
    - triggered_by='file_watcher' or 'manual': an admin already placed
      this file at the live path themselves (dropping it directly into
      storage, or via the admin-side API) -- that action IS the
      approval, so resume_approval_status is set to 'approved'
      immediately and active_storage_key = resume_file_path.

    This distinction exists because a candidate's resume update can
    change more than skills -- employers, dates, titles, project
    descriptions -- all of which feed the tailoring prompt as trusted
    "MASTER_RESUME_TEXT". Gating only the extracted skills (and not the
    file itself) would leave that broader trust gap open."""
    run = ResumeIngestionRun(
        candidate_id=candidate.id,
        resume_file_path=resume_file_path,
        file_hash=file_hash,
        status="processing",
        triggered_by=triggered_by,
    )

    if triggered_by == "candidate_upload":
        run.pending_storage_key = pending_storage_key or resume_file_path
        run.resume_approval_status = "pending"
    else:
        run.pending_storage_key = None
        run.active_storage_key = resume_file_path
        run.resume_approval_status = "approved"

    session.add(run)
    session.flush()

    try:
        text = extract_text_from_docx(file_bytes)
        suggestions = extract_skills_from_resume_text(text)

        existing_skill_names = {
            row.skill_name.strip().lower()
            for row in session.query(SkillInventoryItem)
            .filter_by(candidate_id=candidate.id)
            .all()
        }

        new_count = 0
        for s in suggestions:
            if s.skill_name.strip().lower() in existing_skill_names:
                continue  # already tracked (pending, approved, or rejected) -- don't duplicate
            session.add(
                SkillInventoryItem(
                    candidate_id=candidate.id,
                    skill_name=s.skill_name,
                    tier=s.tier,
                    source_bullet=s.source_bullet,
                    source_project_or_role=s.source_project_or_role,
                    suggested_by="claude_extraction",
                    confidence=s.confidence,
                    status="pending",
                    ingestion_run_id=run.id,
                )
            )
            existing_skill_names.add(s.skill_name.strip().lower())
            new_count += 1

        run.status = "completed"
        run.new_skills_suggested = new_count

    except Exception as e:  # noqa: BLE001 -- deliberately broad: any failure must
        # still leave an auditable failed run row, not a silently half-done ingest.
        run.status = "failed"
        run.error_message = str(e)

    from datetime import datetime, timezone
    run.completed_at = datetime.now(timezone.utc)
    session.commit()
    session.refresh(run)
    return run
